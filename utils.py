"""Utility functions for document loading, vectorstore, analysis, and chat.

This module provides comprehensive utilities for the DocMind AI application,
including document processing, vector database operations, text analysis,
and chat functionality. It supports multiple document formats and provides
advanced features like late chunking, multi-vector embeddings, and hybrid
search capabilities.

Key functionalities:
- Document loading from various formats (PDF, DOCX, CSV, etc.)
- Vector store creation and management with Qdrant
- Document analysis with customizable prompts and chunking
- Chat interface with context retrieval and reranking
- Hardware detection for optimal model suggestions
- Token estimation and text processing utilities

Example:
    Basic document processing workflow::

        docs = load_documents(uploaded_files, late_chunking=True)
        vectorstore = create_vectorstore(docs, multi_vector=True)
        results = analyze_documents(llm, texts, prompt_type, ...)

"""

import logging
import os
import re
import subprocess
import tempfile
from collections.abc import Generator

import extract_msg
import fitz
import nltk
import polars as pl
import tiktoken
import torch
from docx import Document as DocxDocument
from langchain.chains import LLMChain, RetrievalQA
from langchain.chains.summarize import load_summarize_chain
from langchain.document_loaders import TextLoader
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.prompts import PromptTemplate
from langchain.retrievers import ContextualCompressionRetriever
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings.fastembed import FastEmbedSparse
from langchain_community.vectorstores.qdrant import Qdrant
from langchain_core.output_parsers import PydanticOutputParser
from peft import PeftModel
from qdrant_client import QdrantClient
from sentence_transformers import CrossEncoder
from torch import mean
from transformers import AutoModel, AutoTokenizer

from models import AnalysisOutput, Settings
from prompts import INSTRUCTIONS, LENGTHS, PREDEFINED_PROMPTS, TONES

nltk.download("punkt", quiet=True)


def setup_logging():
    """Set up basic logging configuration to file.

    Creates the log directory if it doesn't exist and configures logging
    to write INFO level messages and above to the specified log file with
    timestamp, level, and message formatting.

    Returns:
        None

    Raises:
        OSError: If the log directory cannot be created or the log file
            cannot be written to.
    """
    settings = Settings()
    os.makedirs(os.path.dirname(settings.log_path), exist_ok=True)
    logging.basicConfig(
        filename=settings.log_path,
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )


def detect_hardware() -> tuple[str, int | None]:
    """Detect available hardware and VRAM for optimal model suggestion.

    Attempts to detect NVIDIA GPU availability and VRAM by calling nvidia-smi
    and parsing its output. This information is used to suggest appropriate
    models based on hardware capabilities.

    Returns:
        A tuple containing:
            - Hardware description string ("GPU detected" or "CPU only")
            - VRAM in GB as integer, or None if no GPU detected

    Example:
        >>> hardware, vram = detect_hardware()
        >>> print(f"Hardware: {hardware}, VRAM: {vram}GB")
        Hardware: GPU detected, VRAM: 16GB
    """
    try:
        output = subprocess.check_output(
            ["nvidia-smi"], stderr=subprocess.DEVNULL
        ).decode()
        vram_match = re.search(r"(\d+)MiB / (\d+)MiB", output)
        if vram_match:
            total_vram = int(vram_match.group(2)) // 1024
            return "GPU detected", total_vram
    except (subprocess.CalledProcessError, OSError):
        pass
    return "CPU only", None


def load_pdf_or_epub(file_path: str) -> list[Document]:
    """Load text and images from PDF or EPUB files using PyMuPDF.

    Extracts all text content and converts each page to PNG images for
    multimodal processing. The images are stored in the document metadata.

    Args:
        file_path: Path to the PDF or EPUB file to load.

    Returns:
        List containing a single Document with combined text content and
        page images stored in metadata under the 'images' key.

    Raises:
        fitz.FileDataError: If the file is corrupted or cannot be opened.
        fitz.FileNotFoundError: If the file does not exist.
    """
    doc = fitz.open(file_path)
    text = ""
    images = []
    for page in doc:
        text += page.get_text()
        pix = page.get_pixmap()
        images.append(pix.tobytes("png"))
    return [
        Document(page_content=text, metadata={"source": file_path, "images": images})
    ]


def load_docx_or_odt_or_rtf_or_pptx(file_path: str) -> list[Document]:
    """Load text from DOCX, ODT, RTF, PPTX using python-docx.

    Extracts all paragraph text from Microsoft Office and OpenDocument format
    files by joining paragraphs with newlines.

    Args:
        file_path: Path to the document file to load.

    Returns:
        List containing a single Document with the extracted text content.

    Raises:
        docx.opc.exceptions.PackageNotFoundError: If the file is not a valid
            Office document or cannot be opened.
        FileNotFoundError: If the specified file does not exist.
    """
    doc = DocxDocument(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return [Document(page_content=text, metadata={"source": file_path})]


def load_xlsx_or_csv_or_xml_or_json_or_md(file_path: str) -> list[Document]:
    """Load data from structured files and markdown using Polars or text read.

    Handles multiple structured data formats by using appropriate Polars readers
    for tabular data (Excel, CSV, XML, JSON) and direct text reading for
    markdown files. Tabular data is converted to string representation.

    Args:
        file_path: Path to the file to load. Supported extensions are
            .xlsx, .csv, .xml, .json, and .md.

    Returns:
        List containing a single Document with the file content as text.

    Raises:
        polars.exceptions.ComputeError: If the structured file cannot be parsed.
        FileNotFoundError: If the specified file does not exist.
        UnicodeDecodeError: If markdown file contains invalid encoding.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".xlsx", ".csv", ".xml", ".json"]:
        if ext == ".xlsx":
            df = pl.read_excel(file_path)
        elif ext == ".csv":
            df = pl.read_csv(file_path)
        elif ext == ".xml":
            df = pl.read_xml(file_path)
        elif ext == ".json":
            df = pl.read_json(file_path)
        text = df.to_pandas().to_string()
    elif ext == ".md":
        with open(file_path) as f:
            text = f.read()
    return [Document(page_content=text, metadata={"source": file_path})]


def load_msg(file_path: str) -> list[Document]:
    """Load text from MSG email files using extract-msg.

    Extracts the body text content from Microsoft Outlook MSG email files
    using the extract-msg library.

    Args:
        file_path: Path to the MSG email file to load.

    Returns:
        List containing a single Document with the email body text.

    Raises:
        extract_msg.exceptions.UnsupportedMSGTypeError: If the MSG file type
            is not supported.
        FileNotFoundError: If the specified file does not exist.
    """
    msg = extract_msg.Message(file_path)
    text = msg.body
    return [Document(page_content=text, metadata={"source": file_path})]


def late_chunking(text: str, token_embeddings) -> list:
    """Perform late chunking using NLTK for spans and mean pooling.

    Implements late chunking technique where text is first tokenized into
    sentences, then sentence-level embeddings are computed by averaging
    token embeddings within each sentence span. This approach preserves
    semantic boundaries while enabling efficient processing.

    Args:
        text: Input text to be chunked into sentences.
        token_embeddings: Tensor of token-level embeddings with shape
            (sequence_length, embedding_dim).

    Returns:
        List of sentence-level embeddings as tensors, each representing
        the mean-pooled embedding for one sentence.
    """
    from nltk.tokenize import sent_tokenize

    sentences = sent_tokenize(text)
    spans = []
    start = 0
    for sent in sentences:
        end = start + len(sent)
        spans.append((start, end))
        start = end + 1
    pooled = []
    for start, end in spans:
        chunk_emb = mean(token_embeddings[start:end], dim=0)
        pooled.append(chunk_emb)
    return pooled


def load_documents(uploaded_files, late_chunking=False) -> list[Document]:
    """Load and split documents from uploaded files with optional late chunking.

    Processes multiple file types and converts them to Document objects. Supports
    text extraction from various formats including PDF, Office documents, structured
    data files, and code files. Optionally applies late chunking technique for
    improved embedding quality.

    Args:
        uploaded_files: List of uploaded file objects from Streamlit file uploader.
        late_chunking: Whether to apply late chunking technique using sentence-level
            embedding aggregation for improved semantic representation.

    Returns:
        List of Document objects after text extraction and splitting. Each document
        contains page_content and metadata including source file information.

    Raises:
        ValueError: If an unsupported file type is encountered.
        OSError: If temporary file operations fail.
        RuntimeError: If late chunking model loading fails.
    """
    docs = []
    settings = Settings()
    device = "cuda" if torch.cuda.is_available() else "cpu"
    for file in uploaded_files:
        with tempfile.NamedTemporaryFile(
            delete=False, suffix=os.path.splitext(file.name)[1]
        ) as tmp_file:
            tmp_file.write(file.getvalue())
            file_path = tmp_file.name
        ext = os.path.splitext(file_path)[1].lower()
        if ext in [".pdf", ".epub"]:
            docs.extend(load_pdf_or_epub(file_path))
        elif ext in [".docx", ".odt", ".rtf", ".pptx"]:
            docs.extend(load_docx_or_odt_or_rtf_or_pptx(file_path))
        elif ext in [".xlsx", ".csv", ".xml", ".json", ".md"]:
            docs.extend(load_xlsx_or_csv_or_xml_or_json_or_md(file_path))
        elif ext == ".msg":
            docs.extend(load_msg(file_path))
        elif ext in [".txt", ".py", ".js", ".java", ".ts", ".tsx", ".c", ".cpp", ".h"]:
            loader = TextLoader(file_path)
            docs.extend(loader.load())
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        os.remove(file_path)
    if late_chunking:
        tokenizer = AutoTokenizer.from_pretrained(
            settings.default_embedding_model, trust_remote_code=True
        )
        base_model = AutoModel.from_pretrained(
            settings.default_embedding_model, trust_remote_code=True, device_map="auto"
        )
        model = (
            PeftModel.from_pretrained(base_model, base_model)
            if torch.cuda.is_available()
            else base_model
        )
        for doc in docs:
            inputs = tokenizer(doc.page_content, return_tensors="pt").to(device)
            outputs = model(**inputs)
            token_emb = outputs.last_hidden_state[0]
            chunk_emb = late_chunking(doc.page_content, token_emb)
            doc.metadata["chunk_embeddings"] = chunk_emb
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return splitter.split_documents(docs)


def create_vectorstore(docs: list[Document], multi_vector=False) -> Qdrant:
    """Create Qdrant vectorstore from documents with hybrid search capabilities.

    Initializes a Qdrant vector database with both dense and sparse embeddings
    for hybrid search. Uses HuggingFace embeddings for dense vectors and
    Splade embeddings for sparse vectors to enable both semantic and lexical
    search capabilities.

    Args:
        docs: List of Document objects to be indexed in the vector store.
        multi_vector: Whether to enable multi-vector indexing for improved
            retrieval accuracy with multiple embedding representations.

    Returns:
        Configured Qdrant vectorstore instance ready for similarity search
        and retrieval operations.

    Raises:
        ConnectionError: If unable to connect to the Qdrant server.
        RuntimeError: If embedding model loading fails.
    """
    settings = Settings()
    client = QdrantClient(url=settings.qdrant_url)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    dense_embeddings = HuggingFaceEmbeddings(
        model_name=settings.default_embedding_model,
        model_kwargs={"device_map": "auto" if device == "cuda" else {}},
    )
    sparse_embeddings = FastEmbedSparse(
        model_name="prithivida/Splade_PP_en_v1",
        providers=["CUDAExecutionProvider"] if device == "cuda" else None,
    )
    if multi_vector:
        return Qdrant.from_documents(
            docs,
            dense_embeddings,
            sparse_embedding=sparse_embeddings,
            client=client,
            collection_name="docmind",
            hybrid=True,
            multi_vector=True,
        )
    return Qdrant.from_documents(
        docs,
        dense_embeddings,
        sparse_embedding=sparse_embeddings,
        client=client,
        collection_name="docmind",
        hybrid=True,
    )


def estimate_tokens(text: str) -> int:
    """Estimate token count using tiktoken encoding.

    Uses OpenAI's tiktoken library with the cl100k_base encoding (used by
    GPT-4 and newer models) to provide accurate token count estimates for
    context window management.

    Args:
        text: Input text to count tokens for.

    Returns:
        Estimated number of tokens in the input text.
    """
    encoding = tiktoken.get_encoding("cl100k_base")
    return len(encoding.encode(text))


def aggregate_results(
    results: list[AnalysisOutput | str],
) -> AnalysisOutput | str:
    """Aggregate chunked analysis results into a single output.

    Combines multiple analysis results from document chunks into a unified
    result. Merges summaries, key insights, action items, and open questions
    from all valid AnalysisOutput objects. Falls back to string concatenation
    for non-structured results.

    Args:
        results: List of analysis results, either structured AnalysisOutput
            objects or raw string outputs from failed parsing attempts.

    Returns:
        Single aggregated AnalysisOutput if structured results exist,
        otherwise concatenated string of all results.
    """
    valid_results = [r for r in results if isinstance(r, AnalysisOutput)]
    if not valid_results:
        return " ".join([str(r) for r in results])
    summary = " ".join([r.summary for r in valid_results])
    key_insights = [insight for r in valid_results for insight in r.key_insights]
    action_items = [item for r in valid_results for item in r.action_items]
    open_questions = [q for r in valid_results for q in r.open_questions]
    return AnalysisOutput(
        summary=summary,
        key_insights=key_insights,
        action_items=action_items,
        open_questions=open_questions,
    )


def analyze_documents(
    llm,
    texts: list[str],
    prompt_type: str,
    custom_prompt: str,
    tone: str,
    instruction: str,
    custom_instruction: str,
    length_detail: str,
    context_size: int,
    chunked: bool = False,
) -> list[AnalysisOutput | str]:
    """Analyze documents with customizable prompts and optional chunking.

    Processes documents using the specified LLM with customizable prompts,
    tones, instructions, and length requirements. Handles both single-pass
    and chunked analysis approaches based on document size and user preference.

    Args:
        llm: Language model instance for text generation.
        texts: List of document texts to analyze.
        prompt_type: Type of analysis prompt to use from predefined options.
        custom_prompt: Custom prompt text when prompt_type is "Custom Prompt".
        tone: Desired tone for the analysis output.
        instruction: Type of instructions from predefined options.
        custom_instruction: Custom instruction text when instruction is
            "Custom Instructions".
        length_detail: Desired length/detail level for the output.
        context_size: Maximum context window size for the model.
        chunked: Whether to force chunked analysis regardless of text size.

    Returns:
        List of analysis results, either structured AnalysisOutput objects
        or raw string outputs if parsing fails. Returns single result if
        only one text is analyzed.

    Raises:
        ValueError: If prompt configuration is invalid.
        RuntimeError: If LLM processing fails consistently.
    """
    prompt_text = (
        custom_prompt
        if prompt_type == "Custom Prompt"
        else PREDEFINED_PROMPTS[prompt_type]
    )
    tone_text = TONES[tone]
    instr_text = (
        custom_instruction
        if instruction == "Custom Instructions"
        else INSTRUCTIONS[instruction]
    )
    length_text = LENGTHS[length_detail]

    full_prompt = (
        f"{instr_text} {tone_text} {length_text} {prompt_text} Document: {{text}}"
    )

    parser = PydanticOutputParser(pydantic_object=AnalysisOutput)
    prompt = PromptTemplate(
        template=full_prompt + "\n{format_instructions}",
        input_variables=["text"],
        partial_variables={"format_instructions": parser.get_format_instructions()},
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    results = []
    for text in texts:
        if chunked:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=context_size // 2, chunk_overlap=100
            )
            chunks = splitter.split_text(text)
            chunk_results = []
            for chunk in chunks:
                try:
                    output = chain.run(text=chunk)
                    parsed = parser.parse(output)
                    chunk_results.append(parsed)
                except Exception as e:
                    chunk_results.append(f"Error: {str(e)} - Raw output: {output}")
            aggregated = aggregate_results(chunk_results)
            results.append(aggregated)
        else:
            tokens = estimate_tokens(text)
            if tokens > context_size * 0.8:
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=context_size // 2, chunk_overlap=100
                )
                chunks = splitter.create_documents([text])
                sum_chain = load_summarize_chain(llm, chain_type="map_reduce")
                summary = sum_chain.run(chunks)
                text = summary if isinstance(summary, str) else summary[0].page_content
            try:
                output = chain.run(text=text)
                parsed = parser.parse(output)
                results.append(parsed)
            except Exception as e:
                results.append(f"Error: {str(e)} - Raw output: {output}")
    return results if len(results) > 1 else results[0]


class JinaRerankCompressor:
    """Custom document compressor using local Jina reranker model.

    Uses submodular optimization for document selection and diversity.

    Implements document reranking using CrossEncoder model for relevance scoring
    and submodular selection for diversity optimization. Provides improved
    retrieval quality by balancing relevance and diversity in document selection.

    Attributes:
        model: CrossEncoder instance for document relevance scoring.
        top_n: Maximum number of documents to return after reranking.
    """

    def __init__(self, top_n=5):
        """Initialize the reranker with specified number of top documents.

        Args:
            top_n: Maximum number of documents to return after reranking
                and diversity selection.
        """
        device = "cuda" if torch.cuda.is_available() else "cpu"
        self.model = CrossEncoder(Settings.default_reranker_model, device=device)
        self.top_n = top_n

    def submodular_select(
        self, documents: list[Document], scores: list[float]
    ) -> list[Document]:
        """Simple greedy submodular selection for diversity.

        Uses facility location approximation for optimal selection.

        Implements greedy submodular optimization to select diverse documents
        while maintaining high relevance scores. Uses facility location function
        approximation to balance relevance and diversity.

        Args:
            documents: List of Document objects to select from.
            scores: Relevance scores corresponding to each document.

        Returns:
            List of selected documents optimized for both relevance and diversity,
            limited to self.top_n documents.
        """
        selected = []
        remaining = list(zip(documents, scores, strict=False))
        while len(selected) < self.top_n and remaining:
            best_idx = 0
            best_gain = -float("inf")
            for i, (doc, score) in enumerate(remaining):
                gain = score - max(
                    [
                        self.similarity(doc.page_content, s.page_content)
                        for s in selected
                    ]
                    or [0]
                )
                if gain > best_gain:
                    best_gain = gain
                    best_idx = i
            selected.append(remaining.pop(best_idx)[0])
        return selected

    def similarity(self, text1: str, text2: str) -> float:
        """Calculate text similarity using word overlap (Jaccard similarity).

        Simple similarity metric based on word overlap between two texts.
        In production, this should be replaced with embedding-based similarity
        for better semantic understanding.

        Args:
            text1: First text to compare.
            text2: Second text to compare.

        Returns:
            Similarity score between 0 and 1, where 1 indicates identical
            word sets and 0 indicates no common words.
        """
        return len(set(text1.split()) & set(text2.split())) / len(
            set(text1.split() | text2.split())
        )

    def compress_documents(self, documents, query):
        """Compress and rerank documents based on query relevance.

        Scores all documents against the query using the CrossEncoder model,
        then applies submodular selection to choose diverse, relevant documents.

        Args:
            documents: List of Document objects to compress and rerank.
            query: Query string to score documents against.

        Returns:
            List of top_n documents selected for optimal relevance and diversity.
        """
        pairs = [[query, doc.page_content] for doc in documents]
        scores = self.model.predict(pairs)
        sorted_docs = [
            doc for _, doc in sorted(zip(scores, documents, strict=False), reverse=True)
        ]
        return self.submodular_select(sorted_docs, scores)[: self.top_n]


def chat_with_context(
    llm, vectorstore: Qdrant, user_input: str, history: list[dict]
) -> Generator[str, None, None]:
    """Stream chat response with custom reranker and hybrid retrieval.

    Provides conversational interface with document context retrieval using
    hybrid search and intelligent reranking. Maintains conversation history
    and streams responses for better user experience.

    Args:
        llm: Language model instance for generating responses.
        vectorstore: Qdrant vector database containing indexed documents.
        user_input: Current user question or input.
        history: List of previous conversation turns as dictionaries with
            'user' and 'assistant' keys.

    Yields:
        Streaming text tokens from the language model response.

    Raises:
        ConnectionError: If vector database connection fails.
        RuntimeError: If retrieval or generation fails.
    """
    compressor = JinaRerankCompressor()
    compression_retriever = ContextualCompressionRetriever(
        base_compressor=compressor,
        base_retriever=vectorstore.as_retriever(
            search_type="hybrid", search_kwargs={"k": 10, "score_threshold": 0.5}
        ),
    )
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm, chain_type="stuff", retriever=compression_retriever
    )
    history_str = "\n".join(
        [f"User: {m['user']}\nAssistant: {m['assistant']}" for m in history]
    )
    input_text = f"History: {history_str}\nUser: {user_input}"
    yield from qa_chain.llm.stream(input_text)
